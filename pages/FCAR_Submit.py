# import streamlit as st
# import pandas as pd
# from pathlib import Path
# from datetime import date
# from utils import get_departments, get_courses, upsert_fcar,logo,add_sidebar_logo,summarize_fcar,inject_custom_css
# inject_custom_css()
# st.set_page_config(page_title="FCAR ", page_icon=str(logo), layout="wide")
# # st.set_page_config(page_title="FCAR (Detailed)", page_icon="📝", layout="wide")
# st.markdown("<h1 style='color:#046d5a;'>FCAR Stats</h1>", unsafe_allow_html=True)
# # st.title("Faculty Course Assessment Report (FCAR) — Detailed")
# st.caption("Page for editing and submitting new FCAR entries")
# deps = get_departments()
# courses = get_courses()
# dept_ids = deps["dept_id"].tolist() if not deps.empty else []
# course_ids = courses["course_id"].tolist() if not courses.empty else []
# DATA_PATH = Path(__file__).parents[1] / "data" / "fcar_detail.csv"
# DATA_PATH.parent.mkdir(parents=True, exist_ok=True)
# # # ===== FCAR summary =====
# summary, fcar_df = summarize_fcar()

# k1, k2, k3 = st.columns(3)
# k1.metric("Total FCAR", summary.get("total", 0))
# status_df = summary.get("status", pd.DataFrame(columns=["status", "count"]))
# k2.metric("Statuses", len(status_df))
# by_dept = summary.get("by_dept", pd.DataFrame())
# k3.metric("Active Departments", int(by_dept["dept_id"].nunique()) if not by_dept.empty else 0)

# st.subheader("FCAR Status Overview")
# if not status_df.empty:
#     st.bar_chart(status_df.set_index("status"))
# else:
#     st.info("No FCAR data yet.")

# st.subheader("All FCAR Table")
# st.dataframe(fcar_df,width='stretch')

# st.markdown("---")

# # Quick KPIs
# summary, df = summarize_fcar()
# col1, col2, col3 = st.columns(3)
# col1.metric("Total FCAR Submissions", summary.get("total", 0))
# if "status" in summary:
#     status_df = summary["status"]
# else:
#     status_df = pd.DataFrame(columns=["status","count"])
# col2.metric("Statuses Tracked", len(status_df))
# dept_count = summary.get("by_dept", pd.DataFrame())
# col3.metric("Departments Active", int(dept_count["dept_id"].nunique()) if not dept_count.empty else 0)
# st.subheader("Recent FCAR Entries")
# st.dataframe(df, width='stretch')
import os
import re
import json
from io import StringIO, BytesIO
from pathlib import Path
from datetime import date, datetime
from typing import List, Optional, Dict, Any

import streamlit as st
import pandas as pd

from utils import (
    get_departments,
    get_courses,
    logo,
    add_sidebar_logo,
    summarize_fcar,
    inject_custom_css,
    pct_to_dot,
)
# login guard
if "user" not in st.session_state:
    st.error("You must be logged in to access this page. Please go to the Home page and log in.")
    st.stop()

# Optional: show who is logged in in the sidebar
st.sidebar.markdown(f"👤 **User:** {st.session_state.get('user')} ({st.session_state.get('role')})")

# ---- Page setup ----
inject_custom_css()
# st.set_page_config(page_title="FCAR", page_icon=str(logo), layout="wide")
st.markdown("<h1 style='color:#046d5a;'>Submit / Edit FCAR</h1>", unsafe_allow_html=True)
st.caption("Upload a previous FCAR to auto-fill, or enter data manually. Supports dynamic SLO/PC rows.")

# ---------- Paths ----------
# In Streamlit, __file__ is available when running as a page.
# If you ever run into a context where __file__ isn't set, fall back to CWD.
BASE_DIR = Path(__file__).resolve().parents[1] if "__file__" in globals() else Path.cwd()
DATA_DIR = BASE_DIR / "data"
DATA_DIR.mkdir(parents=True, exist_ok=True)
DETAIL_PATH = DATA_DIR / "fcar_detail.csv"
PCS_PATH = DATA_DIR / "fcar_pcs.csv"
ATTACH_DIR = DATA_DIR / "attachments"
ATTACH_DIR.mkdir(parents=True, exist_ok=True)

# ---------- Lookups ----------
deps = get_departments()
courses = get_courses()
dept_ids = deps["dept_id"].tolist() if isinstance(deps, pd.DataFrame) and not deps.empty else []
course_ids = courses["course_id"].tolist() if isinstance(courses, pd.DataFrame) and not courses.empty else []


# --------- Lazy imports for optional features ---------
def lazy_import_docx():
    try:
        from docx import Document
        return Document
    except Exception:
        st.error("Missing or failed dependency `python-docx`. Install it in your venv:\n\n`pip install python-docx`")
        raise


def lazy_import_pydantic():
    try:
        from pydantic import BaseModel, Field, conint, confloat
        return BaseModel, Field, conint, confloat
    except Exception:
        st.warning("Optional: `pydantic` not installed. JSON validation will be skipped.")
        return None, None, None, None


def lazy_openai_client():
    try:
        from openai import OpenAI
        api_key = os.getenv("OPENAI_API_KEY")
        if not api_key:
            st.error("OPENAI_API_KEY is not set. Add it to your environment to enable AI extraction.")
            return None
        return OpenAI(api_key=api_key)
    except Exception as e:
        st.error(f"OpenAI client not available: {e}")
        return None


# Import for mongodb
def lazy_import_pymongo():
    try:
        import pymongo  # noqa: F401
        from pymongo import MongoClient, ASCENDING
        return pymongo, MongoClient, ASCENDING
    except Exception:
        st.error("Missing dependency `pymongo`. Install it: `pip install pymongo`")
        raise


def get_mongo_client():
    _, MongoClient, _ = lazy_import_pymongo()
    uri = os.getenv("MONGODB_URI") or st.secrets.get("MONGODB_URI", "")
    if not uri:
        st.error("MONGODB_URI not set (env or st.secrets).")
        return None
    try:
        return MongoClient(uri, serverSelectionTimeoutMS=5000)
    except Exception as e:
        st.error(f"Mongo connection failed: {e}")
        return None


def get_db(db_name: str = "fcar_db"):
    client = get_mongo_client()
    return client[db_name] if client else None


def ensure_indexes(db):
    # Unique FCAR id on headers; not unique on pcs since multiple rows per fcar
    _, _, ASCENDING = lazy_import_pymongo()
    db.fcar_headers.create_index([("fcar_id", ASCENDING)], unique=True)
    db.fcar_pcs.create_index([("fcar_id", ASCENDING)], unique=False)
    # Optional: text index for quick search
    db.fcar_headers.create_index(
        [("course_code_title", "text"), ("instructor_name", "text")],
        name="hdr_text_idx",
    )


def save_fcar_to_mongo(header_row: dict, pcs_rows: list, db_name: str = "fcar_db") -> bool:
    """
    Writes:
      - One document to `fcar_headers`
      - N documents to `fcar_pcs`
    If the same fcar_id exists, it will be replaced (upsert).
    """
    db = get_db(db_name)
    if not db:
        raise RuntimeError("No DB connection")

    ensure_indexes(db)

    fcar_id = header_row.get("fcar_id")
    if not fcar_id:
        raise ValueError("header_row must contain fcar_id")

    # Normalize types for safety (strings for text fields)
    header_row = {k: ("" if v is None else v) for k, v in header_row.items()}

    # Upsert header
    db.fcar_headers.update_one(
        {"fcar_id": fcar_id},
        {"$set": header_row},
        upsert=True,
    )

    # Replace PCs for this FCAR (simple approach)
    db.fcar_pcs.delete_many({"fcar_id": fcar_id})
    cleaned = []
    for r in pcs_rows or []:
        cleaned.append(
            {
                "fcar_id": fcar_id,
                "code": str(r.get("code", "")),
                "title": str(r.get("title", "")),
                "method": str(r.get("method", "")),
                "students": int(r.get("students", 0) or 0),
                "avg": float(r.get("avg", 0.0) or 0.0),
                "pct": int(r.get("pct", 0) or 0),
                "standard": str(r.get("standard", "")),
            }
        )
    if cleaned:
        db.fcar_pcs.insert_many(cleaned)

    return True


# --------- Pydantic models (validate LLM output) ---------
BaseModel, Field, conint, confloat = lazy_import_pydantic()


if BaseModel:
    class CDCell(BaseModel):
        strengths_weaknesses: str = ""
        suggestions_improvement: str = ""


    class CourseDelivery(BaseModel):
        teaching_strategies: CDCell = CDCell()
        course_content: CDCell = CDCell()
        resources_facilities: CDCell = CDCell()
        assessment_measures: CDCell = CDCell()


    class PCItem(BaseModel):
        code: str = Field(default="")
        title: str = ""
        method: str = ""
        students: conint(ge=0) = 0
        avg: confloat(ge=0.0) = 0.0
        pct: conint(ge=0, le=100) = 0
        standard: str = ""


    class FcarExtract(BaseModel):
        term: str = ""
        dept_id: Optional[str] = None
        course_id: Optional[str] = None
        course_code_title: str = ""
        section: str = ""
        semester_hours: str = ""
        course_coordinator: str = ""
        academic_year_semester: str = ""
        instructor_name: str = ""
        grades: Dict[str, int] = Field(
            default_factory=lambda: {
                "A": 0,
                "A-": 0,
                "B+": 0,
                "B": 0,
                "B-": 0,
                "C+": 0,
                "C": 0,
                "C-": 0,
                "D+": 0,
                "D": 0,
                "F": 0,
                "I": 0,
                "NP": 0,
                "P": 0,
                "W": 0,
                "Total": 0,
            }
        )
        pcs: List[Dict[str, Any]] = Field(default_factory=list)
        # legacy (kept optional; not shown in UI)
        strengths_weaknesses: str = ""
        suggestions_improvement: str = ""
        teaching_strategies: str = ""
        course_content: str = ""
        resources_facilities: str = ""
        instructor_reflection: str = ""
        # structured 4×2 table
        course_delivery: "CourseDelivery" = CourseDelivery()  # type: ignore[name-defined]


# --------- Word reading ---------
def read_docx(uploaded_file) -> Dict[str, Any]:
    """Return {'text': <all text>, 'tables': [ [ [cell...], ... ], ... ] }"""
    Document = lazy_import_docx()
    doc = Document(uploaded_file)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    text = "\n".join(paragraphs)

    tables = []
    for t in doc.tables:
        t_rows = []
        for r in t.rows:
            t_rows.append([c.text.strip() for c in r.cells])
        tables.append(t_rows)

    return {"text": text, "tables": tables}


# --------- Build the LLM prompt ---------
def build_extraction_prompt(doc_payload: Dict[str, Any]) -> str:
    # Use strings only — NO Python types like int/float here
    schema_hint = {
        "term": "string",
        "dept_id": "string (if present)",
        "course_id": "string (if present)",
        "course_code_title": "string",
        "section": "string",
        "semester_hours": "string",
        "course_coordinator": "string",
        "academic_year_semester": "string",
        "instructor_name": "string",
        "grades": {
            "A": "int",
            "A-": "int",
            "B+": "int",
            "B": "int",
            "B-": "int",
            "C+": "int",
            "C": "int",
            "C-": "int",
            "D+": "int",
            "D": "int",
            "F": "int",
            "W": "int",
            "I": "int",
            "Total": "int",
        },
        "pcs": [
            {
                "code": "string",
                "title": "string",
                "method": "string",
                "students": "int",
                "avg": "float",
                "pct": "int (0-100)",
                "standard": "string",
            }
        ],
        "strengths_weaknesses": "string",
        "suggestions_improvement": "string",
        "teaching_strategies": "string",
        "course_content": "string",
        "resources_facilities": "string",
        "instructor_reflection": "string",
        "course_delivery": {
            "teaching_strategies": {
                "strengths_weaknesses": "string",
                "suggestions_improvement": "string",
            },
            "course_content": {
                "strengths_weaknesses": "string",
                "suggestions_improvement": "string",
            },
            "resources_facilities": {
                "strengths_weaknesses": "string",
                "suggestions_improvement": "string",
            },
            "assessment_measures": {
                "strengths_weaknesses": "string",
                "suggestions_improvement": "string",
            },
        },
    }

    prompt = f"""
You are an expert data extractor. Extract FCAR fields from the provided course report.
Return a SINGLE JSON object that matches exactly this schema (keys must exist; fill missing with reasonable defaults): 
{json.dumps(schema_hint, ensure_ascii=False, indent=2)}

Guidelines:
- Parse grades from FINAL COURSE GRADE DISTRIBUTION table, set non-found grades to empty. Ensure "Total" equals the sum of grade counts if possible.
- Extract all Student Learning Outcomes/Performance Criteria items. Each becomes an entry in "pcs".
- If codes like "PC5.1" are absent, infer short codes ("PC1", "PC2") in reading order.
- Numeric fields: students (int), avg (float), pct (0-100 int). Strip symbols like %.
- Keep text concise; no line breaks inside strings.
- Do NOT include any commentary—ONLY the JSON object.
- Course Delivery 4×2 matrix (REQUIRED): Map table rows to keys:
    Rows (area): “Teaching Strategies”, “Course Content”, “Resources & Facilities”, “Assessment Measures”.
    Columns: “Strengths & Weaknesses”, “Suggestions for Improvement”.
  Normalize header variants (case/spacing/synonyms), e.g.:
    - Teaching Strategies → teaching_strategies (accept “Strategies”, “Instructional Strategies”)
    - Resources & Facilities → resources_facilities (accept “Facilities & Resources”, “Resources/Facilities”)
    - Assessment Measures → assessment_measures (accept “Assessment Methods/Measures”, “Assessment/evidence”)
    - Strengths & Weaknesses (accept “Strengths/Weaknesses”, “Strengths and Weaknesses”)
    - Suggestions for Improvement (accept “Improvements”, “Recommendations”, “Actionable Suggestions”)
- If the document has no explicit table, infer the best cell for each paragraph and still populate the matrix.
- Do not invent content: if a cell is missing, return an empty string for that cell.

--- DOCUMENT TEXT ---
{doc_payload['text']}

--- DOCUMENT TABLES (rows) ---
{json.dumps(doc_payload['tables'], ensure_ascii=False)}
"""
    return prompt


# --------- LLM call ---------
def llm_extract_fields(prompt: str) -> Dict[str, Any]:
    client = lazy_openai_client()
    if client is None:
        raise RuntimeError("AI extraction is not available. Set OPENAI_API_KEY and install openai.")
    model = os.getenv("LLM_MODEL", "gpt-4o-mini")
    resp = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": "You produce strict JSON only."},
            {"role": "user", "content": prompt},
        ],
        temperature=0,
        response_format={"type": "json_object"},
    )
    content = resp.choices[0].message.content
    return json.loads(content)


# ---------- Helpers ----------
GRADE_LABELS = ["A", "A-", "B+", "B", "B-", "C+", "C", "C-", "D+", "D", "F", "W", "I", "Total"]
GRADE_DEFAULTS = {g: 0 for g in GRADE_LABELS}


def num(label, value=0, min_value=0, step=1, key=None):
    return st.number_input(label, min_value=min_value, step=step, value=value, key=key)


def save_attachments(files, fcar_id: str):
    if not files:
        return ""
    dest = ATTACH_DIR / fcar_id
    dest.mkdir(parents=True, exist_ok=True)
    names = []
    for f in files:
        out = dest / f.name
        out.write_bytes(f.read())
        names.append(f.name)
    return "; ".join(names)


def read_any_table(upload):
    """Read CSV / Excel / JSON -> pd.DataFrame or dict. Returns dict with keys:
       header_df, pcs_df, header_record, pcs_list (best-effort)."""
    result = {"header_df": None, "pcs_df": None, "header_record": None, "pcs_list": None}
    if upload is None:
        return result
    name = (upload.name or "").lower()
    try:
        if name.endswith(".csv"):
            df = pd.read_csv(upload)
            result["header_df"] = df
        elif name.endswith(".xlsx") or name.endswith(".xls"):
            xls = pd.ExcelFile(upload)
            prefer = [s for s in xls.sheet_names if s.lower() in ["fcar", "header", "summary", "sheet1"]]
            pcs_sheets = [s for s in xls.sheet_names if s.lower() in ["pcs", "slos", "outcomes"]]
            result["header_df"] = pd.read_excel(xls, prefer[0] if prefer else xls.sheet_names[0])
            if pcs_sheets:
                result["pcs_df"] = pd.read_excel(xls, pcs_sheets[0])
        elif name.endswith(".json"):
            raw = json.load(upload)
            if isinstance(raw, dict):
                if "header" in raw:
                    result["header_record"] = raw.get("header")
                if "pcs" in raw:
                    result["pcs_list"] = raw.get("pcs")
                if "data" in raw and isinstance(raw["data"], list) and raw["data"]:
                    result["header_df"] = pd.DataFrame(raw["data"])
            elif isinstance(raw, list):
                result["header_df"] = pd.DataFrame(raw)
        else:
            try:
                df = pd.read_csv(upload)
                result["header_df"] = df
            except Exception:
                pass
    except Exception as e:
        st.warning(f"Could not parse file: {e}")
    return result


def infer_prefill_from_flat_row(row: pd.Series) -> dict:
    pre = {}
    for k in [
        "course_code_title",
        "course_code",
        "course_title",
        "section",
        "semester_hours",
        "course_coordinator",
        "academic_year_semester",
        "instructor_name",
        "term",
        "dept_id",
        "course_id",
    ]:
        if k in row and pd.notna(row[k]):
            pre[k] = row[k]
    grades = {}
    for g in GRADE_LABELS:
        val = None
        if f"grade_{g}" in row and pd.notna(row.get(f"grade_{g}")):
            val = row[f"grade_{g}"]
        elif g in row and pd.notna(row.get(g)):
            val = row[g]
        grades[g] = int(val) if pd.notna(val) else 0
    pre["grades"] = grades
    for k in [
        "strengths_weaknesses",
        "suggestions_improvement",
        "teaching_strategies",
        "course_content",
        "resources_facilities",
        "instructor_reflection",
    ]:
        if k in row and pd.notna(row[k]):
            pre[k] = row[k]
    pcs = []
    if "pcs_json" in row and pd.notna(row["pcs_json"]):
        try:
            pcs = json.loads(row["pcs_json"])
        except Exception:
            pass
    if not pcs:
        pc_keys: Dict[str, Dict[str, Any]] = {}
        for col in row.index:
            m = re.match(r"^(PC[^_]+)_([a-zA-Z]+)$", str(col))
            if m:
                code, field = m.groups()
                pc_keys.setdefault(code, {})[field] = row[col]
        for code, fields in pc_keys.items():
            pcs.append(
                {
                    "code": code,
                    "title": fields.get("title", ""),
                    "method": fields.get("method", ""),
                    "students": int(fields.get("n", 0) or 0),
                    "avg": float(fields.get("avg", 0.0) or 0.0),
                    "pct": int(fields.get("pct", 0) or 0),
                    "standard": fields.get("standard", fields.get("std", "")),
                }
            )
    pre["pcs"] = pcs
    return pre


def infer_prefill(upload):
    parsed = read_any_table(upload)
    prefill = {"grades": GRADE_DEFAULTS.copy(), "pcs": []}
    if parsed["header_record"]:
        prefill.update(parsed["header_record"])
        if parsed["pcs_list"]:
            prefill["pcs"] = parsed["pcs_list"]
        return prefill
    if isinstance(parsed["header_df"], pd.DataFrame) and not parsed["header_df"].empty:
        row = parsed["header_df"].iloc[0]
        prefill.update(infer_prefill_from_flat_row(row))
    if isinstance(parsed["pcs_df"], pd.DataFrame) and not parsed["pcs_df"].empty:
        cols = {c.lower(): c for c in parsed["pcs_df"].columns}

        def col(name):
            return cols.get(name, name)

        pcs_df = parsed["pcs_df"]
        pcs_rows = []
        for _, r in pcs_df.iterrows():
            pcs_rows.append(
                {
                    "code": str(r.get(col("code"), "")),
                    "title": r.get(col("title"), ""),
                    "method": r.get(col("method"), ""),
                    "students": int(r.get(col("students"), 0) or 0),
                    "avg": float(r.get(col("avg"), 0.0) or 0.0),
                    "pct": int(r.get(col("pct"), 0) or 0),
                    "standard": r.get(col("standard"), ""),
                }
            )
        if not prefill.get("pcs"):
            prefill["pcs"] = pcs_rows
    return prefill


DELIVERY_ROWS = [
    "teaching_strategies",
    "course_content",
    "resources_facilities",
    "assessment_measures",
]


def ensure_delivery_defaults():
    if "course_delivery" not in st.session_state:
        st.session_state["course_delivery"] = {
            k: {"strengths_weaknesses": "", "suggestions_improvement": ""} for k in DELIVERY_ROWS
        }


def ensure_session_defaults():
    if "pcs_rows" not in st.session_state:
        st.session_state["pcs_rows"] = [
            {
                "code": "e.g PC5.1",
                "title": "e.g Apply project management principles",
                "method": "e.g Scoring Rubrics",
                "students": 10,
                "avg": 0.0,
                "pct": 0,
                "standard": "> 3",
            }
        ]
    if "grades" not in st.session_state:
        st.session_state["grades"] = GRADE_DEFAULTS.copy()


def load_prefill_into_session(prefill: dict):
    # Grades
    if "grades" in prefill and isinstance(prefill["grades"], dict):
        st.session_state["grades"] = {g: str(prefill["grades"].get(g, "")) for g in GRADE_LABELS}
        for g in GRADE_LABELS:
            st.session_state[f"grade_{g}"] = str(prefill["grades"].get(g, ""))
    # PCs
    if "pcs" in prefill and isinstance(prefill["pcs"], list) and prefill["pcs"]:
        rows = []
        for r in prefill["pcs"]:
            rows.append(
                {
                    "code": str(r.get("code", "")),
                    "title": r.get("title", ""),
                    "method": r.get("method", ""),
                    "students": int(r.get("students", 0) or 0),
                    "avg": float(r.get("avg", 0.0) or 0.0),
                    "pct": int(r.get("pct", 0) or 0),
                    "standard": r.get("standard", ""),
                }
            )
        st.session_state["pcs_rows"] = rows


# ---------- UI: Mode ----------
left, right = st.columns([2, 1])
with left:
    mode = st.radio("How would you like to start?", ["Manual entry", "Upload Word (AI)"], horizontal=True)
with right:
    word_upload = st.file_uploader("Upload FCAR (.docx)", type=["docx"]) if mode == "Upload Word (AI)" else None

prefill: Dict[str, Any] = {}
ensure_session_defaults()
ensure_delivery_defaults()

# ---------- Upload & AI extraction ----------
if mode == "Upload Word (AI)" and word_upload is not None:
    with st.spinner("Reading Word and extracting fields with AI..."):
        try:
            doc_payload = read_docx(word_upload)
            prompt = build_extraction_prompt(doc_payload)
            raw = llm_extract_fields(prompt)
            if BaseModel:
                # Massage grade keys
                data = FcarExtract(**raw)  # type: ignore
                prefill = data.model_dump()
            else:
                # Best effort when pydantic not installed
                prefill = raw
                prefill.setdefault("grades", GRADE_DEFAULTS.copy())
                prefill.setdefault("pcs", prefill.get("pcs", []) or [])

            # Load into session
            st.session_state["pcs_rows"] = prefill.get("pcs", st.session_state["pcs_rows"])
            st.session_state["grades"] = prefill.get("grades", st.session_state["grades"])
            for g in GRADE_LABELS:
                if "grades" in prefill and isinstance(prefill["grades"], dict):
                    st.session_state[f"grade_{g}"] = str(prefill["grades"].get(g, ""))

            st.success("AI extraction complete. Review and edit below, then Save.")
            with st.expander("Extracted JSON"):
                st.json(prefill)

        except Exception as e:
            st.error(f"Extraction failed: {e}. You can still proceed with Manual entry.")
            prefill = {}
else:
    prefill = {}

# After you build `prefill`
cd_prefill = (prefill or {}).get("course_delivery") or {}
if cd_prefill:
    # validate shape and fill blanks
    out = {k: {"strengths_weaknesses": "", "suggestions_improvement": ""} for k in DELIVERY_ROWS}
    for r in DELIVERY_ROWS:
        cell = cd_prefill.get(r) or {}
        out[r]["strengths_weaknesses"] = str(cell.get("strengths_weaknesses", "") or "")
        out[r]["suggestions_improvement"] = str(cell.get("suggestions_improvement", "") or "")
    st.session_state["course_delivery"] = out

# ========== FORM ==========
with st.form("fcar_form"):
    st.subheader("Course Information")
    c1, c2, c3 = st.columns(3)
    with c1:
        # simple dropdown, default = no selection
        dept_options = [""] + dept_ids if dept_ids else ["(no departments)"]
        dept_id = st.selectbox(
            "Department",
            dept_options,
            index=0,  # default is blank
            format_func=lambda x: "(Select Department)" if x == "" else x,
        )

    with c2:
        course_id = st.text_input(
            "Course Code",
            value=str(prefill.get("course_id", "")),
            placeholder="e.g., CSC490",
        )

    with c3:
        term = st.text_input("Term", value=str(prefill.get("term", "")))

    c4, c5, c6 = st.columns(3)
    with c4:
        course_code = st.text_input(
            "Course Code & Title",
            value=str(prefill.get("course_code_title", "")),
        )
    with c5:
        section = st.text_input("Section(s)", value=str(prefill.get("section", "")))
    with c6:
        sem_hours = st.text_input("Semester Hours", value=str(prefill.get("semester_hours", "")))

    c7, c8, c9 = st.columns(3)
    with c7:
        coordinator = st.text_input("Course Coordinator", value=str(prefill.get("course_coordinator", "")))
    with c8:
        acad_year_sem = st.text_input(
            "Academic Year & Semester",
            value=str(prefill.get("academic_year_semester", "")),
            placeholder="e.g., Spring 2020",
        )
    with c9:
        instructor = st.text_input("Instructor Name", value=str(prefill.get("instructor_name", "")))

    st.markdown("---")
    st.subheader("Final Course Grade Distribution")
    cols = st.columns(8)
    grades: Dict[str, Any] = {}
    for i, g in enumerate(GRADE_LABELS):
        key = f"grade_{g}"
        # ensure session_state key exists (empty by default) so widget uses it
        if key not in st.session_state:
            st.session_state[key] = ""
        with cols[i % 8]:
            # do NOT pass `value=` when also using session_state for the same key
            grades[g] = st.text_input(g, key=key, placeholder="")

    st.caption("Tip: Total should equal the sum of grade counts.")

    st.markdown("---")
    st.subheader("Student Learning Outcomes Assessment (Dynamic)")
    st.caption("Add/edit SLO/PC rows below. Use the + button to add more rows.")
    pcs_df = pd.DataFrame(st.session_state["pcs_rows"])
    left_col_name = " "  # for adding color code
    # pcs_df.insert(0, left_col_name, pcs_df["pct"].apply(pct_to_dot))

    # Only insert the status column once
    if left_col_name not in pcs_df.columns:
     pcs_df.insert(0, left_col_name, pcs_df["pct"].apply(pct_to_dot))
    else:
        pcs_df[left_col_name] = pcs_df["pct"].apply(pct_to_dot)
   
   
    edited_live = st.data_editor(
        pcs_df,
        num_rows="dynamic",
        width="content",
        disabled=(left_col_name,),  # circle column is read-only
        column_config={
            left_col_name: st.column_config.TextColumn("", help="Auto status"),
            "code": st.column_config.TextColumn("Code", help="e.g., PC5.1"),
            "title": st.column_config.TextColumn("Title / Description"),
            "method": st.column_config.TextColumn(
                "Assessment Method",
                help="Rubrics / Exam / Embedded / Survey...",
            ),
            "students": st.column_config.NumberColumn("# Students", min_value=0, step=1),
            "avg": st.column_config.NumberColumn("Average Score", step=0.01, format="%.2f"),
            "pct": st.column_config.NumberColumn("% Met Standard", min_value=0, max_value=100, step=1),
            "standard": st.column_config.TextColumn(
                "Performance Standard",
                help="e.g., > 3 or ≥ 70%",
            ),
        },
        hide_index=True,
    )
    edited_live[left_col_name] = edited_live["pct"].apply(pct_to_dot)
    st.session_state["pcs_rows"] = edited_live.fillna(
        {"code": "", "title": "", "method": "", "students": 0, "avg": 0.0, "pct": 0, "standard": ""}
    ).to_dict(orient="records")

    # --------- Course Delivery (4×2 table) ---------
    st.subheader("Course Delivery")
    cd_state = st.session_state.get("course_delivery", {})
    cd_df = pd.DataFrame.from_dict(cd_state, orient="index")
    cd_df.index.name = "Area"
    cd_df = cd_df.loc[DELIVERY_ROWS]
    cd_df_display = cd_df.rename(
        index={
            "teaching_strategies": "Teaching Strategies",
            "course_content": "Course Content",
            "resources_facilities": "Resources & Facilities",
            "assessment_measures": "Assessment Measures",
        },
        columns={
            "strengths_weaknesses": "Strengths & Weaknesses",
            "suggestions_improvement": "Suggestions for Improvement",
        },
    )

    cd_edited = st.data_editor(
        cd_df_display,
        num_rows="fixed",
        hide_index=False,
        use_container_width=True,
        column_config={
            "Strengths & Weaknesses": st.column_config.TextColumn("Strengths & Weaknesses"),
            "Suggestions for Improvement": st.column_config.TextColumn("Suggestions for Improvement"),
        },
    )

    # back to internal format
    st.session_state["course_delivery"] = cd_edited.rename(
        index={
            "Teaching Strategies": "teaching_strategies",
            "Course Content": "course_content",
            "Resources & Facilities": "resources_facilities",
            "Assessment Measures": "assessment_measures",
        },
        columns={
            "Strengths & Weaknesses": "strengths_weaknesses",
            "Suggestions for Improvement": "suggestions_improvement",
        },
    ).to_dict(orient="index")

    st.subheader("Instructor’s Reflection")
    reflection = st.text_area(
        "Instructor’s Reflection",
        value=str(prefill.get("instructor_reflection", "")),
        height=120,
        placeholder="e.g., No issues to report. All outcomes met.",
    )

    st.markdown("---")
    attachments = st.file_uploader(
        "Attach evidence (rubrics, embedded questions, surveys, etc.)",
        accept_multiple_files=True,
    )

    submitted = st.form_submit_button("Save FCAR")

    if submitted:
        today = date.today().isoformat()
        timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
        safe_dept = (dept_id or "").strip().replace(" ", "_")
        safe_course = (course_id or "").strip().replace(" ", "_")
        safe_sec = (section or "").strip().replace(" ", "_")
        safe_term = (term or "").strip().replace(" ", "_")
        fcar_id = f"{safe_dept}-{safe_course}-{safe_sec}-{safe_term}-{timestamp}"

        attachment_names = save_attachments(attachments, fcar_id)

        def cd_cell(row_key: str, col_key: str) -> str:
            """Safely get a cell from course_delivery table."""
            return (st.session_state.get("course_delivery", {}).get(row_key, {}) or {}).get(col_key, "")

        header_row = {
            "fcar_id": fcar_id,
            "submitted_on": today,
            "dept_id": dept_id,
            "course_id": course_id,
            "course_code_title": course_code,
            "section": section,
            "semester_hours": sem_hours,
            "course_coordinator": coordinator,
            "academic_year_semester": acad_year_sem,
            "instructor_name": instructor,
            "term": term,
            **{f"grade_{g}": grades.get(g, "") for g in GRADE_LABELS},
            # 4×2 matrix flattened
            "cd_teaching_strategies_strengths": cd_cell("teaching_strategies", "strengths_weaknesses"),
            "cd_teaching_strategies_suggestions": cd_cell("teaching_strategies", "suggestions_improvement"),
            "cd_course_content_strengths": cd_cell("course_content", "strengths_weaknesses"),
            "cd_course_content_suggestions": cd_cell("course_content", "suggestions_improvement"),
            "cd_resources_facilities_strengths": cd_cell("resources_facilities", "strengths_weaknesses"),
            "cd_resources_facilities_suggestions": cd_cell("resources_facilities", "suggestions_improvement"),
            "cd_assessment_measures_strengths": cd_cell("assessment_measures", "strengths_weaknesses"),
            "cd_assessment_measures_suggestions": cd_cell("assessment_measures", "suggestions_improvement"),
            "instructor_reflection": reflection,
            "attachments": attachment_names,
            "pcs_json": json.dumps(st.session_state["pcs_rows"]),
        }

        df_new = pd.DataFrame([header_row])
        if DETAIL_PATH.exists():
            try:
                old = pd.read_csv(DETAIL_PATH)
                df_all = pd.concat([old, df_new], ignore_index=True)
            except Exception:
                df_all = df_new
        else:
            df_all = df_new
        df_all.to_csv(DETAIL_PATH, index=False)

        pcs_rows = []
        for r in st.session_state["pcs_rows"]:
            if not any([r.get("code"), r.get("title"), r.get("method")]):
                continue
            pcs_rows.append(
                {
                    "fcar_id": fcar_id,
                    "code": r.get("code", ""),
                    "title": r.get("title", ""),
                    "method": r.get("method", ""),
                    "students": int(r.get("students", 0) or 0),
                    "avg": float(r.get("avg", 0.0) or 0),
                    "pct": int(r.get("pct", 0) or 0),
                    "standard": r.get("standard", ""),
                }
            )

        pcs_df_out = pd.DataFrame(pcs_rows)
        # If you still want the CSV dump, uncomment:
        # if PCS_PATH.exists():
        #     try:
        #         old_pcs = pd.read_csv(PCS_PATH)
        #         pcs_all = pd.concat([old_pcs, pcs_df_out], ignore_index=True)
        #     except Exception:
        #         pcs_all = pcs_df_out
        # else:
        #     pcs_all = pcs_df_out
        # pcs_all.to_csv(PCS_PATH, index=False)

        try:
            ok = save_fcar_to_mongo(header_row, pcs_rows, db_name="fcar_db")
            if ok:
                st.success(
                    "Success! FCAR saved to MongoDB.\n"
                    f"• fcar_id: {fcar_id}\n"
                    "• Header collection: fcar_headers\n"
                    "• PCs collection: fcar_pcs"
                )
                # Show the inserted header for confirmation
                st.dataframe(pd.DataFrame([header_row]), width="content")
        except Exception as e:
            st.error(f"MongoDB save failed: {e}")
